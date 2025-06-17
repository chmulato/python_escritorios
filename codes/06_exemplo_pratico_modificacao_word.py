# ================================================================
# Leitura e Modificação de Arquivo Word - Exemplo Prático
#
# Este script lê um arquivo Word (.docx), verifica a existência de uma cláusula
# de confidencialidade e, caso não exista, adiciona a cláusula ao final do documento.
# A expressão "Cláusula de Confidencialidade:" será impressa em negrito.
#
# Autor: Christian Mulato
# Data: Junho de 2025
#
# ATENÇÃO:
# - Certifique-se de instalar a biblioteca python-docx antes de executar: pip install python-docx
# - Verifique se o arquivo Word está no caminho correto.
# ================================================================
import docx

# Caminhos dos arquivos
input_path = r'C:\dev\python_escritorios\codes\contrato_exemplo.docx'
output_path = r'C:\dev\python_escritorios\codes\contrato_confidencial.docx'

# Texto da cláusula de confidencialidade
clausula = (
    "Cláusula de Confidencialidade: As partes se comprometem a manter em sigilo todas as informações "
    "trocadas em razão deste contrato, não podendo divulgá-las a terceiros sem prévia autorização por escrito."
)

# Abrir o documento
doc = docx.Document(input_path)

# Verificar se a cláusula já existe
texto_doc = "\n".join([p.text for p in doc.paragraphs])
if "Cláusula de Confidencialidade" not in texto_doc:
    # Adicionar parágrafo ao final
    par = doc.add_paragraph()
    run_negrito = par.add_run("Cláusula de Confidencialidade:")
    run_negrito.bold = True
    par.add_run(" As partes se comprometem a manter em sigilo todas as informações trocadas em razão deste contrato, não podendo divulgá-las a terceiros sem prévia autorização por escrito.")

# Salvar o documento modificado
doc.save(output_path)
print("Cláusula de confidencialidade adicionada (se necessário) e documento salvo com sucesso.")